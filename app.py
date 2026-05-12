import streamlit as st
import google.generativeai as genai
import data_helper
from docx import Document
from io import BytesIO

st.set_page_config(page_title="Arquiteto de Ensino UNIPAR", layout="wide")

st.title("🍎 Arquiteto de Ensino Medicina UNIPAR")
st.caption("Foco em DCNs 2025, Matriz ENAMED e Integração Curricular")

# 1. Autenticação e Configuração do Modelo
if "GEMINI_API_KEY" not in st.secrets:
    st.error("ERRO: Configure a 'GEMINI_API_KEY' nos Secrets do Streamlit.")
    st.stop()

genai.configure(api_key=st.secrets["GEMINI_API_KEY"])

try:
    modelos = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
    # Seleção automática do melhor modelo disponível
    target = 'models/gemini-1.5-flash' if 'models/gemini-1.5-flash' in modelos else modelos[0]
    model = genai.GenerativeModel(target)
except Exception as e:
    st.error(f"Erro na conexão: {e}")
    st.stop()

# --- FUNÇÃO PARA GERAR DOCX ---
def gerar_docx(conteudo, titulo_disciplina):
    doc = Document()
    doc.add_heading(f'Plano de Ensino: {titulo_disciplina}', 0)
    # Adiciona o conteúdo gerado pela IA ao documento
    doc.add_paragraph(conteudo)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- INTERFACE ---
with st.sidebar:
    st.header("Identificação")
    professor = st.text_input("Nome do Professor:")
    turma = st.text_input("Turma/Período:")
    st.markdown("---")
    st.caption(f"✅ Motor Ativo: {target}")

st.subheader("1. Configuração do Módulo e Disciplina")
modulo = st.selectbox("Módulo Integrado:", list(data_helper.MODULOS_UNIPAR.keys()))
dados = data_helper.MODULOS_UNIPAR[modulo]
disciplina = st.text_input("Nome da Disciplina Específica:")

st.subheader("2. Objetivos e Temas")
temas = st.text_area("Descreva os temas que você irá ministrar:", height=150)

if st.button("Gerar Plano de Ensino e Habilitar Exportação"):
    if not professor or not disciplina or not temas:
        st.warning("Preencha todos os campos obrigatórios.")
    else:
        prompt = f"""
        Atue como um Especialista em Educação Médica. Crie um Plano de Ensino para a UNIPAR.
        DISCIPLINA: {disciplina} | MÓDULO: {modulo} | PROFESSOR: {professor} | TURMA: {turma}
        CH Módulo: {dados['carga_horaria']} | Ciclo: {dados['ciclo']} | Ementa: {dados['ementas']}
        
        TEMAS DO PROFESSOR: {temas}

        REGRAS:
        1. Objetivos Bloom adequados ao ciclo {dados['ciclo']}.
        2. Tabela de correlação ENAMED (Portaria 478/2025).
        3. Avaliação: 70% Somativa (Provas/OSCE) e 30% Formativa (PBL/Práticas).
        4. Metodologia: Mix integrado e criativo.
        5. Bibliografia ABNT (3 básicas, 3 complementares).
        
        Gere um texto estruturado e profissional.
        """

        with st.spinner('A arquitetar o plano...'):
            try:
                res = model.generate_content(prompt)
                texto_final = res.text
                
                st.markdown("---")
                st.markdown(texto_final)
                
                # Botões de Download
                col1, col2 = st.columns(2)
                
                with col1:
                    st.download_button(
                        label="📥 Baixar em DOCX (Word)",
                        data=gerar_docx(texto_final, disciplina),
                        file_name=f"Plano_{disciplina}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                with col2:
                    st.download_button(
                        label="📄 Baixar em Markdown",
                        data=texto_final,
                        file_name=f"Plano_{disciplina}.md"
                    )
                    
            except Exception as e:
                st.error(f"Erro na geração: {e}")
